from tkinter import *
from pynput import mouse
import csv
import docx
import webbrowser
import tkinter as tk
import tkinter.messagebox
#from tkinter import filedialogs


def openword():
    webbrowser.open("4stepsexam.py")
    webbrowser.open("Answer Sheet.docx")


#     my_program = filedialog.askopenfilename()

#     os.system('"%s"' % my_program)
def btn_clicked():
    if entry0.index("end") == 0:
        validatemail()
        return
    if not entry0.get():
        btn_clicked()

    if entry1.index("end") == 0:
        validatepass()
        return
    if not entry1.get():
        btn_clicked()
    tkinter.messagebox.showinfo("Login", "Login Success, Welcome!")
    b1.place(
        x=766, y=505,
        width=213,
        height=72)


def validatemail():
    tkinter.messagebox.showinfo("Login", "Please enter a valid e-mail")


def validatepass():
    tkinter.messagebox.showinfo("Login", "Please enter a valid password")


window = tk.Tk()

window.geometry("1000x600")
window.configure(bg="#293335")
canvas = tk.Canvas(
    window,
    bg="#293335",
    height=600,
    width=1000,
    bd=0,
    highlightthickness=0,
    relief="ridge")
canvas.place(x=0, y=0)

background_img = tk.PhotoImage(file=f"background.png")
background = canvas.create_image(
    508.5, 228.0,
    image=background_img)

entry0_img = tk.PhotoImage(file=f"img_textBox0.png")
entry0_bg = canvas.create_image(
    166.0, 367.0,
    image=entry0_img)

entry0 = tk.Entry(
    bd=0,
    bg="#ffffff",
    highlightthickness=0)

entry0.place(
    x=22, y=351,
    width=288,
    height=30)

entry1_img = tk.PhotoImage(file=f"img_textBox1.png")
entry1_bg = canvas.create_image(
    166.0, 456.0,
    image=entry1_img)

entry1 = tk.Entry(
    window,
    show="*",
    bd=0,
    bg="#ffffff",
    highlightthickness=0)

entry1.place(
    x=22, y=440,
    width=288,
    height=30)

img0 = tk.PhotoImage(file=f"img0.png")
b0 = tk.Button(
    image=img0,
    borderwidth=0,
    highlightthickness=0,
    command=btn_clicked,
    relief="flat")

b0.place(
    x=28, y=500,
    width=102,
    height=38)
img1 = PhotoImage(file=f"img1.png")

b1 = tk.Button(
    image=img1,
    borderwidth=0,
    highlightthickness=0,
    command=openword,
    relief="flat")

m = open('tst.txt', 'a', newline='')

doc1 = docx.Document("The Expirement wordfile.docx")
doc2 = docx.Document("The expirement Model Answer wordfile.docx")
doc1paragraphs = []
doc2paragraphs = []

for paragraph in doc1.paragraphs:  # We save the paragraphs in lists
    doc1paragraphs.append(paragraph.text)
for paragraph in doc2.paragraphs:
    doc2paragraphs.append(paragraph.text)

for i in doc1paragraphs:  # We check which paragraphs match and which do not
    if i in doc2paragraphs:
        m.write(f"[MATCH   ] {i}")
    else:
        m.write(f"[NO MATCH] {i}")

o = open('dataset.csv', 'a', newline='')

writer = csv.writer(o)
writer.writerow(
    ['Point At X', 'Point At Y', 'Pressed At X', 'Pressed At y', 'Released at X', 'Released at X', 'Scrolled At X',
     'Scrolled At y'])


def on_move(x, y):
    writer.writerow([x, y, 'false', 'false', 'false', 'false', 'false', 'false'])


def on_click(x, y, button, pressed):
    if pressed:
        writer.writerow([x, y, 'true', 'true', 'false', 'false', 'false', 'false']);

    else:
        writer.writerow([x, y, 'false', 'false', 'true', 'true', 'false', 'false']);


def on_scroll(x, y, dx, dy):
    writer.writerow([x, y, 'false', 'false', 'false', 'false', 'true', 'true']);


def main():
    # Collect events until released
    with mouse.Listener(
            on_move=on_move,
            on_click=on_click,
            on_scroll=on_scroll) as listener:
        listener.join()


# ...or, in a non-blocking fashion:
listener = mouse.Listener(
    on_move=on_move,
    on_click=on_click,
    on_scroll=on_scroll)
listener.start()
window.resizable(False, False)
window.mainloop()