import docx
from docx import Document
import os  
import csv  


 #your docx file path
doc=Document('bahgaga2.docx')
o = open('results.txt', 'a', newline='')

for p in doc.paragraphs:
    fonts = []
    texts = []
    sizes = []
    bold= []
    italic=[]
    color=[]

    for i in p.runs:
        texts.extend(i.text.split())
        fonts.append(i.font.name)
        sizes.append(i.font.size.pt)
        bold.append(i.font.bold)
        italic.append(i.font.italic)
        color.append(i.font.color.rgb)

doc=Document('bodda.docx')
for p in doc.paragraphs:
    ifonts = []
    itexts = []
    isizes = []
    ibold= []
    iitalic=[]
    icolor=[]
    for i in p.runs:
        ifonts.append(i.font.name)
        itexts.extend(i.text.split())
        isizes.append(i.font.size.pt)
        ibold.append(i.font.bold)
        iitalic.append(i.font.italic)
        icolor.append(i.font.color.rgb)
    print(texts, sizes,italic,bold,fonts,color, sep="\n")
    print(itexts, isizes,iitalic,ibold,ifonts,icolor, sep="\n")
    
    
def compareSize(l1,l2):
    global Score
    Score=0
    if(l1==l2):
        Score=Score+1
        return "Equal in font size"
    else:
      return "Not equal in font size"
l1=sizes
l2=isizes

def comparefont(l3,l4):
    global Score
    Score=Score
    if(l3==l4):
     Score=Score+1
     return "Equal in font family"
    else:
      return "Not equal in font family"
l3=fonts
l4=ifonts

def comparetext(l5,l6):
    global Score
    Score=Score
    if(l5==l6):
     Score=Score+1
     return "Equal in text"
    else:
      return "Not equal in text"
l5=texts
l6=itexts

def comparebold(l7,l8):
    global Score
    Score=Score
    if(l7==l8):
     Score=Score+1
     return "Equal in bold text"
    else:
      return "Not equal in bold text"
l7=bold
l8=ibold

def compareitalic(l9,l10):
    global Score
    Score=Score
    if(l9==l10):
     Score=Score+1
     return "Equal in italic text"
    else:
        
      return "Not equal in italic text"
l9=italic
l10=iitalic

def comparecolor(l11,l12):
    global Score
    Score=Score
    if(l11==l12):
     Score=Score+1
     return "Equal in text color"
    else:
      return "Not equal in text color"
l11=color
l12=icolor

 
writer = csv.writer(o)
writer.writerow(["comparison",compareSize(l1,l2),"and",comparefont(l3,l4),"and",comparetext(l5,l6),"and",comparebold(l7,l8),"and",compareitalic(l9,l10),"and",comparecolor(l11,l12)," and The Score=",Score])
 
print("comparison",compareSize(l1,l2),"and",comparefont(l3,l4),"and",comparetext(l5,l6),"and",comparebold(l7,l8),"and",compareitalic(l9,l10),"and",comparecolor(l11,l12)," and The Score=",Score) 
