from tkinter import *
from pynput import mouse
import csv
import docx
import webbrowser
import tkinter as tk
import tkinter.messagebox
from tkinter import filedialog
from ctypes import Structure, windll, c_uint, sizeof, byref
import threading
import re 
from db_setup import *
import time

def openword():
    
    webbrowser.open("The Expirement wordfile.docx")
    time.sleep(1.2)
    webbrowser.open("stepsui.py")
    start_listener()
#     my_program = filedialog.askopenfilename()

#     os.system('"%s"' % my_program)
uid = ""
def btn_clicked():
    account = check_login(entry0.get(), entry1.get())
    globals()['uid'] = account[0]
    if (account and account[3] == "user"):
        # tkinter.messagebox.showinfo("login","Vaild email")
        tkinter.messagebox.showinfo("Login","Login Success, Welcome!")
        #btn_clicked()
        
        b1.place(
                x = 766, y = 505,
                width = 213,
                height = 72)
    else:
        tkinter.messagebox.showinfo("login","invalid email or password")

window = tk.Tk()

window.geometry("1000x600")
window.configure(bg = "#293335")
canvas = tk.Canvas(
    window,
    bg = "#293335",
    height = 600,
    width = 1000,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge")
canvas.place(x = 0, y = 0)

background_img = tk.PhotoImage(file = f"background.png")
background = canvas.create_image(
    508.5, 228.0,
    image=background_img)

entry0_img = tk.PhotoImage(file = f"img_textBox0.png")
entry0_bg = canvas.create_image(
    166.0, 367.0,
    image = entry0_img)

entry0 = tk.Entry(
    bd = 0,
    bg = "#ffffff",
    highlightthickness = 0)


entry0.place(
    x = 22, y = 351,
    width = 288,
    height = 30)

entry1_img = tk.PhotoImage(file = f"img_textBox1.png")
entry1_bg = canvas.create_image(
    166.0, 456.0,
    image = entry1_img)

entry1 = tk.Entry(
    window,
    show="*",
    bd = 0,
    bg = "#ffffff",
    highlightthickness = 0)

entry1.place(
    x = 22, y = 440,
    width = 288,
    height = 30)

img0 = tk.PhotoImage(file = f"img0.png")
b0 = tk.Button(
    image = img0,
    borderwidth = 0,
    highlightthickness = 0,
    command = btn_clicked,
    relief = "flat")

b0.place(
    x = 28, y = 500,
    width = 102,
    height = 38)
img1 = PhotoImage(file = f"img1.png")

b1 = tk.Button(
    image = img1,
    borderwidth = 0,
    highlightthickness = 0,
    command = openword,
    relief = "flat")

m = open('tst.txt', 'a', newline='')

doc1 = docx.Document("The Expirement wordfile.docx")
doc2 = docx.Document("The expirement Model Answer wordfile.docx")
doc1paragraphs = []
doc2paragraphs = []
 
for paragraph in doc1.paragraphs: #We save the paragraphs in lists
    doc1paragraphs.append(paragraph.text)
for paragraph in doc2.paragraphs:
    doc2paragraphs.append(paragraph.text)
 
for i in doc1paragraphs: #We check which paragraphs match and which do not
    if i in doc2paragraphs:
         m.write(f"[MATCH   ] {i}")
    else:
         m.write(f"[NO MATCH] {i}")
         
         
o = open('dataset.csv', 'a', newline='')
 
writer = csv.writer(o)
writer.writerow(['Q1'])


def on_move(x, y):

   writer.writerow([str(x)+ ',' + str(y)+ ',false'+ ',false'+ ',false'+ ',false'+ ',false'+ ',false'])
   result = get_idle_duration()
   on_move_db(uid, x, y, result)
    
onClickCtr = 0
def on_click(x, y, button, pressed):
    result = get_idle_duration()
    if pressed:
        globals()['onClickCtr']+=1
        # writer.writerow([str(x)+ ',' +str(y)+ ',true'+ ',true'+ ',false'+ ',false'+ ',false'+ ',false']);
        result = get_idle_duration()
        on_pressed_db(uid, x, y, result)
        on_click_db(uid, x, y, onClickCtr, result)
    else:
        writer.writerow([str(x)+ ',' +str(y)+ ',false'+ ',false'+ ',true'+ ',true'+ ',false'+ ',false']);
        result = get_idle_duration()
        on_released_db(uid, x, y, result)

def on_scroll(x, y, dx, dy):
    writer.writerow([str(x)+',' + str(y)+ ',false'+ ',false'+ ',false'+ ',false'+ ',true'+ ',true']);
    result = get_idle_duration()
    on_scroll_db(uid, x, y, result)  


def printit():
    print(get_idle_duration())
    
class LASTINPUTINFO(Structure):
    _fields_ = [
        ('cbSize', c_uint),
        ('dwTime', c_uint),
    ]

def get_idle_duration():
    threading.Timer(10.0, printit).start()
    lastInputInfo = LASTINPUTINFO()
    lastInputInfo.cbSize = sizeof(lastInputInfo)
    windll.user32.GetLastInputInfo(byref(lastInputInfo))
    millis = windll.kernel32.GetTickCount() - lastInputInfo.dwTime
    return millis / 1000.0

# printit()    

def main():
# Collect events until released
    with mouse.Listener(
        on_move=on_move,
        on_click=on_click,
        on_scroll=on_scroll) as listener:
        listener.join()

# ...or, in a non-blocking fashion
def start_listener():
    listener = mouse.Listener(
    on_move=on_move,
    on_click=on_click,
    on_scroll=on_scroll)
    listener.start()
window.resizable(False, False)
window.mainloop()

