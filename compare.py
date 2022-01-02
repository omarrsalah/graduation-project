import docx
 
doc1 = docx.Document("The expirement Model Answer wordfile.docx")
doc2 = docx.Document("The Expirement wordfile.docx")
doc1paragraphs = []
doc2paragraphs = []
 
for paragraph in doc1.paragraphs: #We save the paragraphs in lists
    doc1paragraphs.append(paragraph.text)
for paragraph in doc2.paragraphs:
    doc2paragraphs.append(paragraph.text)
 
for i in doc1paragraphs: #We check which paragraphs match and which do not
    if i in doc2paragraphs:
         print(f"[MATCH   ] {i}")
    else:
         print(f"[NO MATCH] {i}")